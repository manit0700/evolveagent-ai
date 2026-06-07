import csv
import os
import re
from datetime import UTC, datetime
from pathlib import Path
from uuid import uuid4

from fastapi import UploadFile

from app.config import DATA_DIR
from app.services.storage_service import StorageService


UPLOAD_DIR = Path(DATA_DIR).parent / "uploads"
EXTRACTED_DIR = UPLOAD_DIR / "extracted"
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
MAX_FILES_PER_UPLOAD = 5
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".html",
    ".css",
    ".pdf",
    ".docx",
}
CODE_EXTENSIONS = {".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css"}


class FileService:
    def __init__(self, storage: StorageService):
        self.storage = storage
        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)

    async def process_uploads(
        self,
        files: list[UploadFile],
        session_id: str | None = None,
        workspace_id: str | None = None,
    ) -> list[dict]:
        if len(files) > MAX_FILES_PER_UPLOAD:
            return [
                {
                    "file_id": str(uuid4()),
                    "session_id": session_id,
                    "workspace_id": workspace_id,
                    "filename": "upload batch",
                    "content_type": None,
                    "extension": "",
                    "size_bytes": 0,
                    "status": "failed",
                    "text_preview": "",
                    "extracted_text_length": 0,
                    "error": f"Upload limit is {MAX_FILES_PER_UPLOAD} files per request.",
                }
            ]
        results = []
        for file in files:
            results.append(await self.process_file(file, session_id=session_id, workspace_id=workspace_id))
        return results

    async def process_file(
        self,
        file: UploadFile,
        session_id: str | None = None,
        workspace_id: str | None = None,
    ) -> dict:
        file_id = str(uuid4())
        original_filename = file.filename or "uploaded-file"
        safe_filename = self.safe_filename(original_filename)
        extension = Path(safe_filename).suffix.lower()
        content = await file.read()
        size_bytes = len(content)
        now = datetime.now(UTC).isoformat()
        stored_path = UPLOAD_DIR / f"{file_id}_{safe_filename}"
        extracted_path = EXTRACTED_DIR / f"{file_id}.txt"

        metadata = {
            "file_id": file_id,
            "session_id": session_id,
            "workspace_id": workspace_id,
            "filename": original_filename,
            "safe_filename": safe_filename,
            "stored_path": str(stored_path),
            "extracted_text_path": str(extracted_path),
            "content_type": file.content_type,
            "extension": extension,
            "size_bytes": size_bytes,
            "status": "failed",
            "extracted_text_length": 0,
            "text_preview": "",
            "created_at": now,
            "error": None,
        }

        if extension not in SUPPORTED_EXTENSIONS:
            metadata["error"] = f"Unsupported file type '{extension or 'unknown'}'."
            self.storage.append("files.json", metadata)
            return self.response_metadata(metadata)

        if size_bytes > MAX_FILE_SIZE_BYTES:
            metadata["error"] = "File exceeds the 10 MB limit."
            self.storage.append("files.json", metadata)
            return self.response_metadata(metadata)

        try:
            stored_path.write_bytes(content)
            extracted_text = self.extract_text(stored_path, extension)
            extracted_path.write_text(extracted_text, encoding="utf-8")
            metadata["status"] = "processed"
            metadata["extracted_text_length"] = len(extracted_text)
            metadata["text_preview"] = self.preview(extracted_text)
        except Exception as exc:
            metadata["error"] = f"Could not process file: {exc}"

        self.storage.append("files.json", metadata)
        return self.response_metadata(metadata)

    @staticmethod
    def response_metadata(metadata: dict) -> dict:
        return {
            "file_id": metadata["file_id"],
            "session_id": metadata.get("session_id"),
            "workspace_id": metadata.get("workspace_id"),
            "filename": metadata["filename"],
            "content_type": metadata.get("content_type"),
            "extension": metadata["extension"],
            "size_bytes": metadata["size_bytes"],
            "status": metadata["status"],
            "text_preview": metadata.get("text_preview", ""),
            "extracted_text_length": metadata.get("extracted_text_length", 0),
            "error": metadata.get("error"),
        }

    @staticmethod
    def safe_filename(filename: str) -> str:
        name = os.path.basename(filename).strip() or "uploaded-file"
        name = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
        return name[:120] or "uploaded-file"

    def extract_text(self, path: Path, extension: str) -> str:
        if extension in {".txt", ".md", ".json", ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css"}:
            return self.read_text_file(path)
        if extension == ".csv":
            return self.extract_csv(path)
        if extension == ".pdf":
            return self.extract_pdf(path)
        if extension == ".docx":
            return self.extract_docx(path)
        return ""

    @staticmethod
    def read_text_file(path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def extract_csv(path: Path) -> str:
        try:
            with path.open("r", encoding="utf-8", errors="replace", newline="") as file:
                reader = csv.reader(file)
                rows = list(reader)
        except Exception as exc:
            return f"CSV extraction failed: {exc}"
        if not rows:
            return "CSV file is empty."
        headers = rows[0]
        preview_rows = rows[1:6]
        lines = [
            "CSV data summary:",
            f"- Columns ({len(headers)}): {', '.join(headers)}",
            f"- Row count excluding header: {max(len(rows) - 1, 0)}",
            "- First rows:",
        ]
        for index, row in enumerate(preview_rows, start=1):
            lines.append(f"  {index}. {dict(zip(headers, row)) if headers else row}")
        return "\n".join(lines)

    @staticmethod
    def extract_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text_parts = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(f"--- Page {index} ---\n{text.strip()}")
        extracted = "\n\n".join(text_parts).strip()
        return extracted or "No extractable text found. Scanned PDFs/OCR are not supported in MVP v1.5."

    @staticmethod
    def extract_docx(path: Path) -> str:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    @staticmethod
    def preview(text: str, limit: int = 360) -> str:
        compact = " ".join(text.split())
        return compact[:limit] + ("..." if len(compact) > limit else "")

    def get_file(self, file_id: str) -> dict | None:
        return next((item for item in self.storage.read_list("files.json") if item.get("file_id") == file_id), None)

    def read_extracted_text(self, file_id: str) -> str:
        metadata = self.get_file(file_id)
        if not metadata or metadata.get("status") != "processed":
            return ""
        path = Path(metadata.get("extracted_text_path", ""))
        if not path.exists():
            return ""
        return path.read_text(encoding="utf-8", errors="replace")

    def build_context(self, file_ids: list[str], limit: int = 20_000) -> tuple[str, list[dict]]:
        context_parts = []
        files_used = []
        remaining = limit
        for file_id in file_ids:
            metadata = self.get_file(file_id)
            if not metadata or metadata.get("status") != "processed":
                continue
            text = self.read_extracted_text(file_id)
            if not text:
                continue
            clipped = text[:remaining]
            if not clipped:
                break
            context_parts.append(
                f"File: {metadata.get('filename')} ({metadata.get('extension')})\n"
                f"Extracted text:\n{clipped}"
            )
            remaining -= len(clipped)
            files_used.append(
                {
                    "file_id": metadata["file_id"],
                    "workspace_id": metadata.get("workspace_id"),
                    "filename": metadata["filename"],
                    "content_type": metadata.get("content_type"),
                    "extension": metadata["extension"],
                    "size_bytes": metadata["size_bytes"],
                    "extracted_text_length": metadata.get("extracted_text_length", 0),
                }
            )
            if remaining <= 0:
                break
        return "\n\n---\n\n".join(context_parts), files_used
